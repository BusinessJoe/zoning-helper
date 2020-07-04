import os
import json
import docx
import unidecode

import convert.parsedocx as parsedocx


def is_blank(string):
    return string.isspace() or len(string) == 0


class Paragraph:
    richtext_format = {
        'bold': '<b>{text}</b>',
        'underline': '<u>{text}</u>',
        'italic': '<i>{text}</i>',
        'superscript': '<sup>{text}</sup>',
        'subscript': '<sub>{text}</sub>'
    }

    def __init__(self, para):
        self.paragraph = para

    @property
    def text(self):
        return self.paragraph.text

    @property
    def richtext(self):
        all_text = ''
        for r in self.paragraph.runs:
            text = r.text

            if is_blank(text):
                continue

            if r.bold:
                text = self.richtext_format['bold'].format(text=text)
            if r.underline:
                text = self.richtext_format['underline'].format(text=text)
            if r.font.superscript:
                text = self.richtext_format['superscript'].format(text=text)
            if r.font.subscript:
                text = self.richtext_format['subscript'].format(text=text)

            all_text += text
        return all_text

    @property
    def is_context(self):
        if self.text.isspace() or len(self.text) == 0:
            return False

        para = self.paragraph
        is_bold = all(run.bold for run in para.runs if not run.text.isspace())
        is_underline = all(run.underline for run in para.runs if not run.text.isspace())
        is_uppercase = all(run.text == run.text.upper() for run in para.runs if not run.text.isspace())

        return is_bold and is_underline and is_uppercase


class DocxBylawReader:
    """Reads bylaw specifications and exceptions from docx files"""
    def __init__(self, filename, area):
        self.doc = docx.Document(filename)
        self.area = area
        self.paragraphs = [Paragraph(para) for para in self.doc.paragraphs]

    @property
    def html_text(self):
        all_text = '<br>\n'.join(para.richtext for para in self.paragraphs if not is_blank(para.richtext))
        return unidecode.unidecode(all_text)

    def save_bylaws(self, bylaw_type, path):
        parse = {'spec': parsedocx.parse_specifications,
                 'exc' : parsedocx.parse_exceptions}[bylaw_type]

        os.makedirs(os.path.join(path, self.area), exist_ok=True)
        for bylaw in parse(self.html_text):
            bylaw['area'] = self.area
            filename = os.path.join(path, self.area, f"{bylaw['code']}.json")
            with open(filename, 'w') as f:
                json.dump(bylaw, f)


if __name__ == '__main__':
    reader = DocxBylawReader('convert/CCREST.docx', 'cliffcrest')
    reader.save_bylaws('spec', 'app/static/bylaws/specifications')
    reader.save_bylaws('exc', 'app/static/bylaws/exceptions')

